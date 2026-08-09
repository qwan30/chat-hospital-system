"""DOCX document loader using python-docx."""

from __future__ import annotations

from pathlib import Path
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from hospital_ai.core.errors import ExternalServiceError
from hospital_ai.services.loaders.base import BaseDocumentLoader, LoadedPage


class DocxLoader(BaseDocumentLoader):
    """Extract text from DOCX files.

    Treats each paragraph as content. Header/table content is also extracted.
    """

    def supported_extensions(self) -> set[str]:
        return {".docx", ".doc"}

    def load(self, file_path: Path, mime_type: str = "") -> list[LoadedPage]:
        if not file_path.exists():
            raise ExternalServiceError(f"DOCX file not found: {file_path}")

        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError:
            return self._load_ooxml_without_python_docx(file_path)

        try:
            doc = Document(str(file_path))
        except Exception as exc:
            raise ExternalServiceError(f"Failed to open DOCX: {exc}") from exc

        text_parts: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)
        if not full_text.strip():
            raise ExternalServiceError("DOCX document is empty or contains no extractable text.")

        return [LoadedPage(page_number=1, text=full_text, confidence=1.0)]

    @staticmethod
    def _load_ooxml_without_python_docx(file_path: Path) -> list[LoadedPage]:
        """Read paragraph text from a DOCX OOXML package when python-docx is absent."""
        try:
            with ZipFile(file_path) as archive:
                document_xml = archive.read("word/document.xml")
            root = ElementTree.fromstring(document_xml)
        except (BadZipFile, KeyError, ElementTree.ParseError) as exc:
            raise ExternalServiceError(f"Failed to open DOCX: {exc}") from exc

        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        paragraphs = []
        for paragraph in root.iter(f"{namespace}p"):
            text = "".join(node.text or "" for node in paragraph.iter(f"{namespace}t")).strip()
            if text:
                paragraphs.append(text)

        full_text = "\n".join(paragraphs)
        if not full_text:
            raise ExternalServiceError("DOCX document is empty or contains no extractable text.")
        return [LoadedPage(page_number=1, text=full_text, confidence=1.0)]
